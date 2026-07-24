"""
parser.py - Resume File Parser

Handles extracting raw text from uploaded PDF and DOCX files.
Provides a single public function `extract_text()` that auto-detects
the file type and delegates to the correct parser.
"""

import io  # In-memory binary streams for file handling
from typing import Optional

# Third-party libraries for document parsing
from PyPDF2 import PdfReader       # Reads PDF files
from docx import Document          # Reads DOCX (Word) files


def extract_text(file_bytes: bytes, filename: str) -> Optional[str]:
    """
    Detect the file type from the filename extension and extract
    all readable text from the document.

    Args:
        file_bytes: Raw binary content of the uploaded file.
        filename: Original filename (used to detect .pdf / .docx).

    Returns:
        A single string containing all extracted text, or None
        if the file type is unsupported.
    """

    # Lower-case the filename so the check is case-insensitive
    lower_name = filename.lower()

    # --- PDF extraction ---
    if lower_name.endswith(".pdf"):
        return _extract_from_pdf(file_bytes)

    # --- DOCX extraction ---
    if lower_name.endswith(".docx"):
        return _extract_from_docx(file_bytes)

    # Unsupported file type
    return None


def _extract_from_pdf(file_bytes: bytes) -> str:
    """
    Read a PDF from raw bytes and concatenate text from every page.

    PyPDF2 opens the PDF via a BytesIO stream so we never need to
    touch the filesystem.
    """
    # Wrap raw bytes in a seekable in-memory stream
    stream = io.BytesIO(file_bytes)

    # PdfReader parses the PDF structure
    reader = PdfReader(stream)

    # Collect text from each page
    pages = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            pages.append(page_text)

    # Join all pages with newlines to preserve spacing
    return "\n".join(pages)


def _extract_from_docx(file_bytes: bytes) -> str:
    """
    Read a DOCX file from raw bytes and concatenate text from
    every paragraph.

    python-docx also needs a seekable stream, so we wrap the bytes
    in a BytesIO object.
    """
    stream = io.BytesIO(file_bytes)

    # Document() opens the .docx archive and reads its XML content
    doc = Document(stream)

    # Each paragraph in the document becomes one line of output
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    return "\n".join(paragraphs)
